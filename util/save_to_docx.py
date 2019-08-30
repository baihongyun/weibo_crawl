from docx import Document
from docx.shared import Inches
from io import BytesIO
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
import docx.image


def save_to_docx(content_and_comment_dict):
    #如果没有内容和评论，直接返回
    if(len(content_and_comment_dict) < 0):
        return
    #将内容和评论直接写到文档中
    else:
        document = Document(docx = 'test.docx')
        for key in content_and_comment_dict:
             if key == "time":
                p = document.add_paragraph()
                run = p.add_run(content_and_comment_dict[key])
                run.font.size = Pt(12)
                run.bold = True
             if key == "content":
                p = document.add_paragraph()
                run = p.add_run(content_and_comment_dict[key])
                run.font.size = Pt(12)
             elif key == "comment":
                 pass
                 '''
                  num_comment = len(content_and_comment_dict[key])
                  if num_comment > 0:
                        for content in content_and_comment_dict[key]:
                            document.add_paragraph(content)
                '''
             elif key == "url":
                pass
             elif key == "pictures":
                 num_pictures = len(content_and_comment_dict[key])
                 if num_pictures > 0:
                     table = document.add_table(rows = 1,cols = 2)
                     hdr_cells = table.rows[0].cells
                     index = 0
                     for picture in content_and_comment_dict[key]:
                         if index == 2:
                             index = 0
                             hdr_cells = table.add_row().cells
                         paragraph = hdr_cells[index].paragraphs[0]
                         try:
                             run = paragraph.add_run()
                             run.add_picture(BytesIO(picture), width=Inches(3.0))
                             index = index + 1
                         except docx.image.exceptions.UnrecognizedImageError as e:
                             pass
             else:
                pass
        document.styles['Normal'].font.name = '宋体'  # 设置字体
        document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        document.save('test.docx')
#写封面
def write_cover(user_main_info):
    if len(user_main_info) < 0:
        return
    else:
        document = Document(docx = 'test.docx')
        for value in user_main_info:
            if value == "user_name":
                p = document.add_paragraph()
                run = p.add_run(user_main_info[value])
                run.font.size = Pt(16)
                run.bold = True
            if value == "touxiang_img":
                document.add_picture(BytesIO(user_main_info[value]), width=Inches(2.5))  # 添加图片
            if value == "user_brief":
                p = document.add_paragraph()
                run = p.add_run("简介："+user_main_info[value])
                run.font.size = Pt(12)  # 二号
                run.bold = True
            if value == "weibo_num":
                p = document.add_paragraph()
                run = p.add_run("微博总数："+user_main_info[value])
                run.font.size = Pt(12)  # 二号
                run.bold = True
            if value == "gz_num":
                p = document.add_paragraph()
                run = p.add_run("关注数："+user_main_info[value])
                run.font.size = Pt(12)  # 二号
                run.bold = True
            if value == "fs_num":
                p = document.add_paragraph()
                run = p.add_run("粉丝数："+user_main_info[value])
                run.font.size = Pt(12)  # 二号
                run.bold = True
        document.styles['Normal'].font.name = '宋体'  # 设置字体
        document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        document.save('test.docx')

if __name__ == "__main__":
    content = {'content':'微博内容','comment':['第一条评论','第二条评论','第三条评论']}
    save_to_docx(content)
